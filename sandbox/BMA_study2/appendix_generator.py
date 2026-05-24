import io
import os
import pandas as pd
import docx
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt

# Import your existing pipeline tools
import sys
trgt_dict = os.path.abspath(r'../cbm_prep')
sys.path.append(trgt_dict)
import plotting
from regressions import regression_comp
from objects import RegressionResult


def add_figure_caption(doc, text):
    """
    Helper function to insert a native MS Word dynamic caption.
    This enables cross-referencing and Tables of Figures.
    """
    # Attempt to use the standard 'Caption' style; fallback to normal if not in template
    try:
        p = doc.add_paragraph(style='Caption')
    except KeyError:
        p = doc.add_paragraph()

    p.add_run("Figure ")

    # Create the XML element for a simple field (SEQ Figure \* ARABIC)
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), ' SEQ Figure \\* ARABIC ')

    # Add a fallback/placeholder number '1' to the field
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '1'
    r.append(t)
    fldSimple.append(r)

    # Append the field XML to the paragraph XML
    p._p.append(fldSimple)

    # Add the descriptive text after the dynamic number
    p.add_run(f": {text}")
    return p

def create_word_appendix(
    methd_comp,
    mk_csv_path,
    output_filename="results/Appendix_Regressions.docx",
    ordered_variables=None,
    fig_title_mapping=None,
    doc_title_mapping=None
):
    """
    Generates an MS Word appendix with cell type regressions and MK regression.
    Filters/orders figures based on `ordered_variables` and renames them using mapping (in the figure and for the text in the doc).
    """
    # Setup defaults
    if ordered_variables is None:
        # Fallback to everything in the order it was calculated
        ordered_variables = [rec["variable"] for rec in methd_comp.results]

    if fig_title_mapping is None:
        fig_title_mapping = {}
    if doc_title_mapping is None:
        doc_title_mapping = fig_title_mapping

    # Initialize Word Document
    doc = docx.Document()
    # doc.add_heading('Appendix: Regression Figures', 0)

    fig_num = 1
    base_style = {
        'equal_limits': True,
        'grid': True,
        'ci': True,
        'ci_mode': 'shade'
    }

    # 1. Plot Cell Types from MethodComparator based on ordered_variables
    print("Generating cell type regression figures...")
    for var_name in ordered_variables:
        # Search the MethodComparator results for a matching record
        matching_recs = [rec for rec in methd_comp.results if rec["variable"] == var_name]

        if not matching_recs:
            print(f"Warning: No regression results found for '{var_name}'. Skipping.")
            continue

        rec = matching_recs[0]  # Take the first match

        # Get the display title, fallback to internal var_name if not in mapping
        fig_display_title = fig_title_mapping.get(var_name, var_name)
        doc_display_title = doc_title_mapping.get(var_name, var_name)

        # Merge dynamic labels into the base style
        plot_style = base_style.copy()
        plot_style.update({
            'title': fig_display_title,
            'xlabel': 'Manual Review [%]',
            'ylabel': 'Digital Review [%]'
        })

        # Generate the figure
        fig, ax = plt.subplots(figsize=(6, 5))
        fig, ax = plotting.plot_scatter_basic(rec, style=plot_style, fig=fig, ax=ax)
        plotting.overlay_regression_line(data=rec, fig=fig, ax=ax, style=plot_style)

        # Save figure to a memory buffer
        mem_stream = io.BytesIO()
        fig.savefig(mem_stream, format='png', bbox_inches='tight')
        mem_stream.seek(0)
        plt.close(fig)

        # Add a page break for every figure after the first one
        if fig_num > 1:
            doc.add_page_break()

        # Embed in Word Document with native MS Word caption
        caption = f"Method comparison Deming regression analysis for {doc_display_title}."
        add_figure_caption(doc, caption)
        doc.add_picture(mem_stream, width=Inches(5))
        fig_num += 1

    # 2. Plot MK Regression from standalone CSV
    print("Generating Megakaryocyte regression figure...")
    if os.path.exists(mk_csv_path):
        mk_df = pd.read_csv(mk_csv_path)

        # Run regression calculation
        reg_dict = regression_comp(mk_df['Mean Digital'], mk_df['Mean DSS'], reg_method='deming')
        mk_reg = RegressionResult.from_dict(reg_dict)

        # Package into the mock record dictionary
        mk_rec = {
            "variable": "Megakaryocytes",
            "x": mk_df['Mean Digital'],
            "y": mk_df['Mean DSS'],
            "reg": mk_reg
        }

        fig_display_title = fig_title_mapping.get("Megakaryocytes", "Megakaryocytes")
        doc_display_title = doc_title_mapping.get("Megakaryocytes", "Megakaryocytes")

        mk_style = base_style.copy()
        mk_style.update({
            'title': fig_display_title,
            'xlabel': "Users' Count [MK / 10X FOVs]",
            'ylabel': "DSS Suggestion [MK / 10X FOVs]"
        })

        fig, ax = plt.subplots(figsize=(6, 5))
        fig, ax = plotting.plot_scatter_basic(mk_rec, style=mk_style, fig=fig, ax=ax)
        plotting.overlay_regression_line(data=mk_rec, fig=fig, ax=ax, style=mk_style)

        mem_stream = io.BytesIO()
        fig.savefig(mem_stream, format='png', bbox_inches='tight')
        mem_stream.seek(0)
        plt.close(fig)

        if fig_num > 1:
            doc.add_page_break()

        caption = f"Method comparison Deming regression analysis for {doc_display_title}."
        add_figure_caption(doc, caption)
        doc.add_picture(mem_stream, width=Inches(5))
    else:
        print(f"Warning: Could not find MK raw data at {mk_csv_path}. Skipping MK plot.")

    # Save the document
    os.makedirs(os.path.dirname(output_filename), exist_ok=True)
    doc.save(output_filename)
    print(f"Appendix successfully saved to {output_filename}")
