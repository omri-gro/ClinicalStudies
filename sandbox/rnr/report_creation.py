import pandas as pd
import numpy as np
import warnings
from statsmodels.tools.sm_exceptions import ConvergenceWarning
import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from processing import AC_CONFIG

warnings.simplefilter('ignore', ConvergenceWarning)

# --- CONFIGURATION ---
OUTPUT_DIR = "results"
PLOT_DIR = os.path.join(OUTPUT_DIR, "plots")
os.makedirs(PLOT_DIR, exist_ok=True)

_caption_counters = {'Table': 0, 'Figure': 0}

def plot_precision_profiles(data, study_type="Reproducibility"):
    df = pd.DataFrame(data)
    df = df[df['constant'] == False]
    plot_paths = {}

    for param in df['Parameter'].unique():
        param_data = df[df['Parameter'] == param]
        if param not in AC_CONFIG: continue

        threshold, sd_limit, cv_limit, _ = AC_CONFIG[param]
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

        below_th = param_data[param_data['Mean'] <= threshold]
        above_th = param_data[param_data['Mean'] > threshold]

        if above_th.empty:
            # If all points are below threshold, only scale the axis to the max data point
            max_x = param_data['Mean'].max() * 1.15
            if np.isnan(max_x) or max_x == 0:
                max_x = threshold * 0.5  # Fallback if all values are literally 0
        else:
            # If there are points above threshold, ensure both threshold and max data point are visible
            max_x = max(param_data['Mean'].max() * 1.15, threshold * 1.15)

        # Plot 1: SD vs Mean
        sd_line_end = threshold if not above_th.empty else max_x
        if not below_th.empty:
            ax1.scatter(below_th['Mean'], below_th['Total_SD'], color='#1f77b4', edgecolor='k', s=55, zorder=4, label='Active (Evaluated by SD)')
        if not above_th.empty:
            ax1.scatter(above_th['Mean'], above_th['Total_SD'], color='#a6a6a6', edgecolor='k', s=55, alpha=0.25, zorder=3, label='Muted (Evaluated by %CV)')
            # Only draw vertical line if the threshold is actively being crossed
            ax1.axvline(x=threshold, color='#7f7f7f', linestyle=':', linewidth=1.5, zorder=2)

        ax1.plot([0, sd_line_end], [sd_limit, sd_limit], color='#d62728', linestyle='-', linewidth=2.5, zorder=5,
                 label=f'SD Limit (≤ {sd_limit})')
        ax1.fill_between([0, sd_line_end], 0, sd_limit, color='#2ca02c', alpha=0.07, zorder=1)
        ax1.set_title(f'{param} - {study_type} SD', fontsize=11, fontweight='bold')
        ax1.set_xlabel('Mean Concentration (%)')
        ax1.set_ylabel('Total SD')
        ax1.set_xlim(0, max_x)
        ax1.set_ylim(0, max(sd_limit * 1.6, (below_th['Total_SD'].max() if not below_th.empty else 0) * 1.3, 0.5))
        ax1.legend(loc='upper right', fontsize=8.5)
        ax1.grid(True, linestyle='--', alpha=0.4)

        # Plot 2: CV vs Mean
        if not below_th.empty:
            ax2.scatter(below_th['Mean'], below_th['Total_CV'], color='#a6a6a6', edgecolor='k', s=55, alpha=0.25, zorder=3, label='Muted (Evaluated by SD)')
        if not above_th.empty:
            ax2.scatter(above_th['Mean'], above_th['Total_CV'], color='#2ca02c', edgecolor='k', s=55, zorder=4, label='Active (Evaluated by %CV)')
            ax2.plot([threshold, max_x], [cv_limit, cv_limit], color='#d62728', linestyle='-', linewidth=2.5, zorder=5, label=f'CV Limit (≤ {cv_limit}%)')
            ax2.fill_between([threshold, max_x], 0, cv_limit, color='#2ca02c', alpha=0.07, zorder=1)
            ax2.axvline(x=threshold, color='#7f7f7f', linestyle=':', linewidth=1.5, zorder=2)

        ax2.set_title(f'{param} - {study_type} %CV', fontsize=11, fontweight='bold')
        ax2.set_xlabel('Mean Concentration (%)')
        ax2.set_ylabel('Total %CV')
        ax2.set_xlim(0, max_x)
        ax2.set_ylim(0, max(cv_limit * 2.0, (above_th['Total_CV'].max() if not above_th.empty else 0) * 1.3, 60.0))
        ax2.legend(loc='upper right', fontsize=8.5)
        ax2.grid(True, linestyle='--', alpha=0.4)

        plt.tight_layout()
        img_path = os.path.join(PLOT_DIR, f"{param.replace(' ', '_')}_{study_type}.png")
        plt.savefig(img_path, dpi=150)
        plt.close()
        plot_paths[param] = img_path
    return plot_paths

def set_repeat_table_header(row):
    """ Helper to force MS Word to repeat table headers across pages """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def set_landscape(doc):
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Decrease margins to give tables more horizontal space
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    return section


def add_native_caption(doc, text, prefix="Table"):
    """Injects a native MS Word SEQ field with a cached value for instant rendering."""
    # Increment the counter for this specific prefix (Table or Figure)
    _caption_counters[prefix] = _caption_counters.get(prefix, 0) + 1
    current_num = str(_caption_counters[prefix])

    p = doc.add_paragraph(style='Caption')
    p.add_run(f'{prefix} ')

    # Build the complex field structure Word expects
    run = p.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' SEQ {prefix} \\* ARABIC '
    run._r.append(instrText)

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar_sep)

    # Inject the Python-tracked number as the visible text (cached value)
    t = OxmlElement('w:t')
    t.text = current_num
    run._r.append(t)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

    p.add_run(f'. {text}')


def apply_strict_widths(table, widths_in_inches):
    """Instantly forces MS Word to respect strict column widths and unlocks manual resizing."""

    total_width_twips = sum(int(w * 1440) for w in widths_in_inches)
    table.autofit = False

    # 1. Set Master Table Width
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(total_width_twips))

    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    # 2. THE FIX: Completely clear the corrupted grid and rebuild it flawlessly
    tblGrid = table._tbl.tblGrid
    tblGrid.clear() # Wipe all remaining/corrupted background columns

    for w in widths_in_inches:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * 1440)))
        tblGrid.append(gridCol)

    # 3. Apply exact widths to the individual cells
    for row_idx, tr in enumerate(table._tbl.tr_lst):
        if row_idx < 2:
            # First 2 rows: handle the complex header merges
            row_cells = table.rows[row_idx].cells
            processed_cells = set()
            for col_idx, cell in enumerate(row_cells):
                if cell._tc in processed_cells: continue
                processed_cells.add(cell._tc)

                span = sum(1 for c in row_cells if c._tc == cell._tc)
                width_twips = sum(int(widths_in_inches[col_idx + i] * 1440) for i in range(span) if (col_idx + i) < len(widths_in_inches))

                tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tcW.w = width_twips
                tcW.type = 'dxa'
        else:
            # Data rows: O(1) pure XML injection to maintain blazing fast speed
            col_idx = 0
            for tc in tr.tc_lst:
                tcPr = tc.get_or_add_tcPr()
                gridSpan = tcPr.find(qn('w:gridSpan'))
                span = int(gridSpan.get(qn('w:val'))) if gridSpan is not None else 1

                width_twips = sum(int(widths_in_inches[col_idx + i] * 1440) for i in range(span) if (col_idx + i) < len(widths_in_inches))

                tcW = tcPr.get_or_add_tcW()
                tcW.w = width_twips
                tcW.type = 'dxa'

                col_idx += span

def build_variance_table(doc, data, study_type):
    if not data: return
    total_rows = len(data) + 2
    table = doc.add_table(rows=total_rows, cols=13)
    table.style = 'Table Grid'
    table.autofit = False

    # Repeat first two rows across pages
    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    h1 = table.rows[0].cells
    h2 = table.rows[1].cells

    # Merge vertical headers
    for i, t in enumerate(['Parameter', 'Sample', 'N', 'Mean']):
        h1[i].text = t
        h1[i].merge(h2[i])
    h1[12].text = 'Status'
    h1[12].merge(h2[12])

    h1[4].text = 'Repeatability'
    h1[4].merge(h1[5])
    if study_type == 'Repeatability':
        h1[6].text = 'Between-Run'
        h1[6].merge(h1[7])
        h1[8].text = 'Between-Day'
        h1[8].merge(h1[9])
        h1[10].text = 'Within-Lab (Total)'
        h1[10].merge(h1[11])
    else:
        h1[6].text = 'Between-Day'
        h1[6].merge(h1[7])
        h1[8].text = 'Between-Site'
        h1[8].merge(h1[9])
        h1[10].text = 'Reproducibility (Total)'
        h1[10].merge(h1[11])

    for i in range(4, 12, 2):
        h2[i].text = 'SD'
        h2[i + 1].text = '%CV'

    # Dynamically fetch the _Cell constructor to avoid importing internal classes
    CellClass = type(table.cell(0, 0))
    table_rows = table.rows

    for r_idx, row in enumerate(data):
        tr = table_rows[r_idx + 2]._tr  # Get raw XML row
        tcs = tr.tc_lst  # Get raw XML cells list

        # Wrap XML directly without triggering the grid evaluation
        CellClass(tcs[0], table).text = row['Parameter']
        CellClass(tcs[1], table).text = str(row['Sample'])

        c_n = CellClass(tcs[2], table)
        c_n.text = str(row['N'])
        c_n.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        c_mean = CellClass(tcs[3], table)
        c_mean.text = f"{row['Mean']:.2f}"
        c_mean.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if row.get('constant'):
            # Instant Native XML merge (Bypasses .merge() layout recalculation)
            c_merged = CellClass(tcs[4], table)
            c_merged.text = "Dependent variable is constant."
            c_merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            tcPr = tcs[4].get_or_add_tcPr()
            gridSpan = OxmlElement('w:gridSpan')
            gridSpan.set(qn('w:val'), "8")
            tcPr.append(gridSpan)

            # Remove the absorbed cells from the XML
            for tc in tcs[5:12]:
                tr.remove(tc)

            CellClass(tcs[12], table).text = "N/A"
        else:
            vals = [
                f"{row['Rep_SD']:.2f}",
                f"{row['Rep_CV']:.1f}%",
                f"{row.get('BR_SD', row.get('BD_SD')):.2f}",
                f"{row.get('BR_CV', row.get('BD_CV')):.1f}%",
                f"{row.get('BD_SD', row.get('BS_SD')):.2f}",
                f"{row.get('BD_CV', row.get('BS_CV')):.1f}%",
                f"{row['Total_SD']:.2f}",
                f"{row['Total_CV']:.1f}%",
                row['Status']
            ]
            for i, val in enumerate(vals):
                c = CellClass(tcs[i + 4], table)
                c.text = val
                if i < 8:  # Right-align all numeric columns
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Apply strict widths to all columns
    widths = [2.0, 0.6, 0.35, 0.5, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.6]
    apply_strict_widths(table, widths)


def build_ci_table(doc, data, study_type):
    if not data: return
    doc.add_heading(f'{study_type} - Degrees of Freedom & 95% Confidence Intervals', level=3)

    total_rows = len(data) + 2
    table = doc.add_table(rows=total_rows, cols=10)
    table.style = 'Table Grid'
    table.autofit = False

    # Repeat first two rows across pages
    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    h1 = table.rows[0].cells
    h2 = table.rows[1].cells

    # Merge vertical headers
    for i, t in enumerate(['Parameter', 'Sample', 'N', 'Mean']):
        h1[i].text = t
        h1[i].merge(h2[i])

    h1[4].text = 'DF'
    h1[4].merge(h2[4])
    h1[5].text = 'Repeatability 95% CI'
    h1[5].merge(h1[6])
    h1[7].text = 'DF'
    h1[7].merge(h2[7])
    h1[8].text = f'{"Within-Lab" if study_type == "Repeatability" else "Reproducibility"} 95% CI'
    h1[8].merge(h1[9])

    h2[5].text = 'SD CI'
    h2[6].text = '%CV CI'
    h2[8].text = 'SD CI'
    h2[9].text = '%CV CI'

    CellClass = type(table.cell(0, 0))
    table_rows = table.rows

    for r_idx, row in enumerate(data):
        tr = table_rows[r_idx + 2]._tr
        tcs = tr.tc_lst

        CellClass(tcs[0], table).text = row['Parameter']
        CellClass(tcs[1], table).text = str(row['Sample'])

        for i, key in enumerate(['N', 'Mean']):
            c = CellClass(tcs[i + 2], table)
            c.text = str(row[key]) if key == 'N' else f"{row[key]:.2f}"
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if row.get('constant'):
            c_merged = CellClass(tcs[4], table)
            c_merged.text = "Dependent variable is constant."
            c_merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            tcPr = tcs[4].get_or_add_tcPr()
            gridSpan = OxmlElement('w:gridSpan')
            gridSpan.set(qn('w:val'), "6")  # Span the remaining 6 columns (DFs and CIs)
            tcPr.append(gridSpan)

            # Remove the absorbed cells (indices 5 through 9) from the XML
            for tc in tcs[5:10]:
                tr.remove(tc)
        else:
            vals = [
                str(row['DF_Rep']), row['Rep_SD_CI'], row['Rep_CV_CI'],
                str(row['DF_Total']), row['Total_SD_CI'], row['Total_CV_CI']
            ]
            for i, val in enumerate(vals):
                c = CellClass(tcs[i + 4], table)
                c.text = val
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Apply strict widths to all columns
    # Param, Sample, N, Mean, DF, CI SD, CI CV, DF, CI SD, CI CV
    widths = [2.0, 0.6, 0.35, 0.5, 0.4, 1.25, 1.25, 0.6, 1.25, 1.25]
    apply_strict_widths(table, widths)

def generate_docx(rep_data, repro_data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(8)

    doc.add_heading('Results', level=1)

    # --- CONSOLIDATED TABLES (LANDSCAPE) ---
    set_landscape(doc)

    if rep_data:
        doc.add_heading('1. Repeatability Data Tables', level=2)
        add_native_caption(doc, "Repeatability Variance Components", "Table")
        build_variance_table(doc, rep_data, 'Repeatability')
        doc.add_paragraph("")  # Spacing

        doc.add_page_break()
        add_native_caption(doc, "Repeatability Degrees of Freedom & 95% Confidence Intervals", "Table")
        build_ci_table(doc, rep_data, 'Repeatability')
        doc.add_page_break()

    if repro_data:
        doc.add_heading('2. Reproducibility Data Tables', level=2)
        add_native_caption(doc, "Reproducibility Variance Components", "Table")
        build_variance_table(doc, repro_data, 'Reproducibility')
        doc.add_paragraph("")  # Spacing

        doc.add_page_break()
        add_native_caption(doc, "Reproducibility Degrees of Freedom & 95% Confidence Intervals", "Table")
        build_ci_table(doc, repro_data, 'Reproducibility')
        doc.add_page_break()

    # --- PRECISION PROFILES (PORTRAIT) ---
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading('3. Precision Profiles', level=2)
    doc.add_paragraph(
        "The plots below display Total Precision (SD and %CV) mapped against the Mean for each sample. Acceptance Criteria limits are denoted by the dashed lines.")

    if rep_data:
        doc.add_heading('3.1 Repeatability Profiles', level=3)
        rep_plots = plot_precision_profiles(rep_data, "Repeatability")
        for param, img_path in rep_plots.items():
            add_native_caption(doc, f"{param} Repeatability Profile", "Figure")
            doc.add_picture(img_path, width=Inches(6.0))

    if repro_data:
        doc.add_heading('3.2 Reproducibility Profiles', level=3)
        repro_plots = plot_precision_profiles(repro_data, "Reproducibility")
        for param, img_path in repro_plots.items():
            add_native_caption(doc, f"{param} Reproducibility Profile", "Figure")
            doc.add_picture(img_path, width=Inches(6.0))

    doc.save(os.path.join(OUTPUT_DIR, "RnR_Consolidated_Report.docx"))

